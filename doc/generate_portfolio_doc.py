#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Portfolio Website DOC 생성 스크립트
index.html의 내용을 기반으로 모던한 Word 문서를 생성합니다.

사용 방법:
1. python-docx 설치: pip install python-docx
2. 스크립트 실행: python doc/generate_portfolio_doc.py
   또는 doc 폴더에서: python generate_portfolio_doc.py
3. 생성된 파일: doc/PORTFOLIO_PRESENTATION.docx
"""

import os
import sys
from pathlib import Path

# 현재 스크립트의 디렉토리를 기준으로 경로 설정
SCRIPT_DIR = Path(__file__).parent.absolute()
DOC_DIR = SCRIPT_DIR
ROOT_DIR = SCRIPT_DIR.parent

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def create_portfolio_doc():
    """포트폴리오 DOC 생성"""
    doc = Document()
    
    # 문서 스타일 설정
    style = doc.styles['Normal']
    font = style.font
    font.name = '맑은 고딕'
    font.size = Pt(11)
    
    # 모던한 색상 팔레트
    primary_color = RGBColor(0, 51, 102)  # 진한 파란색
    secondary_color = RGBColor(70, 130, 180)  # 스틸 블루
    accent_color = RGBColor(255, 140, 0)  # 다크 오렌지
    text_color = RGBColor(51, 51, 51)  # 다크 그레이
    
    # 제목: Portfolio
    title = doc.add_heading('Portfolio', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.size = Pt(28)
    title_run.font.bold = True
    title_run.font.color.rgb = primary_color
    
    # 부제목
    subtitle = doc.add_paragraph('Full-Stack Developer')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(16)
    subtitle_run.font.color.rgb = text_color
    
    tech_subtitle = doc.add_paragraph('Java • Spring Framework • Flutter • Mobile Development')
    tech_subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tech_subtitle_run = tech_subtitle.runs[0]
    tech_subtitle_run.font.size = Pt(12)
    tech_subtitle_run.font.color.rgb = secondary_color
    
    doc.add_paragraph()  # 빈 줄
    
    # 1. About Me 섹션
    doc.add_heading('About Me', 1)
    about_heading_run = doc.paragraphs[-1].runs[0]
    about_heading_run.font.color.rgb = primary_color
    
    doc.add_paragraph('Full-Stack 개발자', style='Intense Quote')
    
    doc.add_paragraph('전문 분야:', style='List Bullet')
    doc.add_paragraph('• Java, Spring Framework, Flutter 등 다양한 기술 스택을 활용한 웹 및 모바일 애플리케이션 개발', style='List Bullet 2')
    doc.add_paragraph('• 우리은행, 신한은행, KB국민카드 등 금융권 프로젝트 경험', style='List Bullet 2')
    doc.add_paragraph('• 안드로이드 네이티브 앱 개발부터 백엔드 서버 개발까지 전반적인 개발 역량', style='List Bullet 2')
    
    doc.add_paragraph('주요 통계:', style='List Bullet')
    doc.add_paragraph('✓ 15+ 프로젝트 완료', style='List Bullet 2')
    doc.add_paragraph('✓ 6+ 년 프리랜서 경력', style='List Bullet 2')
    doc.add_paragraph('✓ 4개 자격증 보유', style='List Bullet 2')
    
    doc.add_paragraph('지속적인 학습과 성장을 통해 더 나은 개발자가 되기 위해 노력하고 있으며, 최근에는 Spring Framework 기반 Java Full-Stack 개발자 양성과정을 수료하여 최신 기술을 습득했습니다.')
    
    doc.add_page_break()
    
    # 2. Technical Skills 섹션
    doc.add_heading('Technical Skills', 1)
    skills_heading_run = doc.paragraphs[-1].runs[0]
    skills_heading_run.font.color.rgb = primary_color
    
    # Backend
    doc.add_heading('Backend', 2)
    backend_skills = [
        ('Java', '90%', '객체지향 프로그래밍, 멀티스레딩'),
        ('Spring Framework', '90%', 'MVC, Security, Data JPA'),
        ('Spring Boot', '85%', '마이크로서비스 아키텍처'),
        ('Spring AI', '80%', 'AI 통합 및 LLM 연동'),
        ('JSP/Servlet', '85%', '웹 애플리케이션 개발'),
        ('MyBatis', '85%', '데이터베이스 매핑'),
        ('Python', '80%', '스크립팅 및 자동화')
    ]
    
    for skill, level, desc in backend_skills:
        p = doc.add_paragraph(f'• {skill} ({level}) - {desc}', style='List Bullet')
    
    doc.add_paragraph('주요 경험:', style='List Bullet')
    doc.add_paragraph('✓ RESTful API 설계 및 구현', style='List Bullet 2')
    doc.add_paragraph('✓ Spring Security 기반 인증/인가 시스템', style='List Bullet 2')
    doc.add_paragraph('✓ OAuth2 소셜 로그인 통합', style='List Bullet 2')
    doc.add_paragraph('✓ Spring AI를 활용한 AI 기능 구현', style='List Bullet 2')
    
    # Frontend & Mobile
    doc.add_heading('Frontend & Mobile', 2)
    frontend_skills = [
        ('HTML5/CSS3', '90%', '반응형 웹 디자인'),
        ('Bootstrap', '85%', 'UI 프레임워크'),
        ('JavaScript/jQuery', '85%', '동적 웹 개발'),
        ('Flutter/Dart', '85%', '크로스 플랫폼 개발'),
        ('Android/Java & Kotlin', '90%', '네이티브 앱 개발'),
        ('iOS/Swift & SwiftUI', '80%', 'iOS 앱 개발'),
        ('Python', '80%', '모바일 자동화')
    ]
    
    for skill, level, desc in frontend_skills:
        p = doc.add_paragraph(f'• {skill} ({level}) - {desc}', style='List Bullet')
    
    doc.add_paragraph('주요 경험:', style='List Bullet')
    doc.add_paragraph('✓ Flutter 기반 크로스 플랫폼 앱 개발', style='List Bullet 2')
    doc.add_paragraph('✓ Android 네이티브 앱 개발 (금융권 프로젝트)', style='List Bullet 2')
    doc.add_paragraph('✓ iOS 앱 개발 및 배포', style='List Bullet 2')
    doc.add_paragraph('✓ 반응형 웹 애플리케이션 개발', style='List Bullet 2')
    
    # Database & Tools
    doc.add_heading('Database & Tools', 2)
    db_tools = [
        ('Oracle', '85%', '엔터프라이즈 데이터베이스'),
        ('Git/GitHub & GitLab & Bitbucket', '85%', '버전 관리'),
        ('CI/CD (Jenkins)', '75%', '지속적 통합/배포'),
        ('Docker', '80%', '컨테이너화'),
        ('Figma', '80%', 'UI/UX 디자인')
    ]
    
    for tool, level, desc in db_tools:
        p = doc.add_paragraph(f'• {tool} ({level}) - {desc}', style='List Bullet')
    
    doc.add_paragraph('주요 경험:', style='List Bullet')
    doc.add_paragraph('✓ Oracle 23 AI 데이터베이스 설계 및 최적화', style='List Bullet 2')
    doc.add_paragraph('✓ Git 기반 협업 및 코드 리뷰', style='List Bullet 2')
    doc.add_paragraph('✓ Docker를 활용한 컨테이너화 및 배포', style='List Bullet 2')
    doc.add_paragraph('✓ CI/CD 파이프라인 구축 및 관리', style='List Bullet 2')
    
    doc.add_page_break()
    
    # 3. Key Experience 섹션
    doc.add_heading('Key Experience', 1)
    exp_heading_run = doc.paragraphs[-1].runs[0]
    exp_heading_run.font.color.rgb = primary_color
    
    experiences = [
        {
            'company': '우리은행',
            'project': 'WON뱅킹 Re-Modeling',
            'period': '2022.07 - 2023.07 (12개월)',
            'description': [
                '우리은행 개인비대면 채널 Re-Modeling 추진사업',
                '만보기 기능 추가',
                '이체기능 네이티브 → 웹 서비스 전환',
                '로컬 CI/CD 환경 구축'
            ],
            'tech': 'Android, Java, Kotlin, WebView, CI/CD'
        },
        {
            'company': '신한은행',
            'project': '땡겨요 O2O 플랫폼',
            'period': '2021.10 - 2022.02 (5개월)',
            'description': [
                '음식주문중개 O2O 플랫폼 구축',
                'Pull refresh 확장기능 개발',
                '땡기기 기능 구현',
                'WebView 설계 및 구현',
                'Docker를 이용한 암호화/빌드 시스템 관리'
            ],
            'tech': 'Android, Java, Docker, WebView'
        },
        {
            'company': 'KB 국민카드',
            'project': 'MyData 플랫폼',
            'period': '2021.04 - 2021.08 (5개월)',
            'description': [
                'KB 국민카드 표준API기반 MyData 플랫폼 개편 프로젝트',
                '표준API기반 MyData 기능 적용',
                '전체메뉴 > 메뉴검색 기능 추가'
            ],
            'tech': 'Android, Java, RESTful API'
        },
        {
            'company': '하나은행',
            'project': 'Line Bank Indonesia',
            'period': '2020.08 - 2021.03 (8개월)',
            'description': [
                '인도네시아 하나은행 Linebank 앱 개발',
                'MVVM 패턴 설계 및 구현',
                '보안 키패드 이슈 해결',
                'Django & Bootstrap을 활용한 내부용 앱 배포 사이트 구축'
            ],
            'tech': 'Android, Java, MVVM, Django, Bootstrap'
        },
        {
            'company': 'KB국민은행',
            'project': '마이머니 App 고도화',
            'period': '2019.08 - 2019.11 (4개월)',
            'description': [
                'KB국민은행 마이머니 Android App 고도화 작업',
                '안드로이드 네이티브 앱 개발',
                '인트로 화면/프로그레스바 고도화',
                '지문인증 솔루션 업데이트',
                'androidX 컨버팅'
            ],
            'tech': 'Android, Java, AndroidX'
        },
        {
            'company': '키움증권',
            'project': '영웅문S MTS 개발',
            'period': '2018.05 - 2018.12 (8개월)',
            'description': [
                '키움증권 영웅문S MTS 고도화 프로젝트',
                '관심종목 C++ 공통 플랫폼 개발',
                'Javascript를 이용한 MTS 화면개발'
            ],
            'tech': 'C++, JavaScript, WebView'
        }
    ]
    
    for exp in experiences:
        doc.add_heading(f"{exp['company']} - {exp['project']}", 2)
        p = doc.add_paragraph(exp['period'])
        p_run = p.runs[0]
        p_run.font.bold = True
        p_run.font.color.rgb = secondary_color
        
        doc.add_paragraph('프로젝트 내용:', style='List Bullet')
        for desc in exp['description']:
            doc.add_paragraph(f'• {desc}', style='List Bullet 2')
        
        tech_p = doc.add_paragraph(f"기술 스택: {exp['tech']}")
        tech_p_run = tech_p.runs[0]
        tech_p_run.font.italic = True
        tech_p_run.font.color.rgb = accent_color
        
        doc.add_paragraph()  # 빈 줄
    
    doc.add_page_break()
    
    # 4. Featured Projects 섹션
    doc.add_heading('Featured Projects', 1)
    proj_heading_run = doc.paragraphs[-1].runs[0]
    proj_heading_run.font.color.rgb = primary_color
    
    # Miracle Reading System
    doc.add_heading('Miracle Reading System', 2)
    doc.add_paragraph('프로젝트 개요: 독서 습관 형성과 도서 관리를 위한 종합적인 웹 애플리케이션')
    
    period_p = doc.add_paragraph('개발 기간: 2025.11.10 - 2025.12.10 (1개월)')
    period_p_run = period_p.runs[0]
    period_p_run.font.bold = True
    
    doc.add_paragraph('개발 형태: 개인 프로젝트 (1인 총괄 개발)')
    
    doc.add_paragraph('주요 기능:', style='List Bullet')
    mrs_features = [
        '사용자 인증 및 관리 (폼 로그인, OAuth2)',
        '도서 관리 시스템 (알라딘 API 연동)',
        'AI 기반 도서 요약 (Spring AI + Ollama)',
        '독서 계획 및 기록 관리',
        '속독 훈련 기능',
        '갤러리 및 소셜 기능 (좋아요, 댓글)',
        '마인드맵 기능',
        '관리자 콘솔'
    ]
    for feature in mrs_features:
        doc.add_paragraph(f'✓ {feature}', style='List Bullet 2')
    
    tech_p = doc.add_paragraph('기술 스택: Java 17, Spring Boot 3.3.5, Spring AI, Oracle 23 AI, JSP, Bootstrap 5, jQuery, Docker, Ollama (Qwen3:1.7b)')
    tech_p_run = tech_p.runs[0]
    tech_p_run.font.italic = True
    tech_p_run.font.color.rgb = accent_color
    
    doc.add_paragraph('주요 성과:', style='List Bullet')
    doc.add_paragraph('• AI 통합: Spring AI를 활용한 로컬 LLM 연동', style='List Bullet 2')
    doc.add_paragraph('• 확장 가능한 아키텍처: 계층형 구조 설계', style='List Bullet 2')
    doc.add_paragraph('• 다중 인증 시스템: 폼 로그인 + OAuth2 통합', style='List Bullet 2')
    
    doc.add_paragraph()  # 빈 줄
    
    # Productivity Hub
    doc.add_heading('Productivity Hub', 2)
    doc.add_paragraph('프로젝트 개요: Flutter 기반의 통합 생산성 앱')
    
    period_p = doc.add_paragraph('개발 기간: 2025.12.04 오후 (4시간)')
    period_p_run = period_p.runs[0]
    period_p_run.font.bold = True
    
    doc.add_paragraph('개발 형태: 개인 프로젝트 (1인 총괄 개발)')
    
    doc.add_paragraph('주요 기능:', style='List Bullet')
    ph_features = [
        '할 일 관리 (Todo) - 추가/수정/삭제, 완료 상태 토글',
        '아이디어 기록 - 카테고리별 아이디어 관리',
        '독서 카드 - 독서 진행 관리, 키워드/요약 기록',
        '날씨 정보 - 현재 위치 및 도시별 날씨 조회',
        '뉴스 피드 - AI/양자컴퓨팅 관련 최신 뉴스'
    ]
    for feature in ph_features:
        doc.add_paragraph(f'✓ {feature}', style='List Bullet 2')
    
    tech_p = doc.add_paragraph('기술 스택: Flutter 3.x, Dart, Provider, SQLite, Open-Meteo API, RSS Feed, Geolocator')
    tech_p_run = tech_p.runs[0]
    tech_p_run.font.italic = True
    tech_p_run.font.color.rgb = accent_color
    
    doc.add_paragraph('아키텍처:', style='List Bullet')
    doc.add_paragraph('• Provider 패턴 (MVVM 기반) 상태 관리', style='List Bullet 2')
    doc.add_paragraph('• SQLite 로컬 데이터 저장', style='List Bullet 2')
    doc.add_paragraph('• RESTful API 연동 (날씨, 뉴스)', style='List Bullet 2')
    doc.add_paragraph('• 반응형 UI 디자인', style='List Bullet 2')
    
    doc.add_page_break()
    
    # 5. Education & Certifications 섹션
    doc.add_heading('Education & Certifications', 1)
    edu_heading_run = doc.paragraphs[-1].runs[0]
    edu_heading_run.font.color.rgb = primary_color
    
    doc.add_heading('학위', 2)
    doc.add_paragraph('• 석사 - 광주과학기술원 기전공학과 (2005.03 ~ 2007.08)', style='List Bullet')
    doc.add_paragraph('• 학사 - 강원대학교 전기전자공학과 (1995.03 ~ 2004.02)', style='List Bullet')
    
    doc.add_heading('교육 이력', 2)
    educations = [
        ('Spring Framework 기반 Java Full-Stack 개발자 양성과정', '쌍용강북교육센터', '2025.05.12 - 2025.11.12', '944시간'),
        ('소음진동평가모니터링시스템개발 과정', '경영기술개발원교육센터', '2012.06 - 2012.12', '960시간'),
        ('임베디드 SW 전문가 과정', '한국정보기술연구원', '2007.10 - 2008.03', '960시간')
    ]
    
    for edu, org, period, hours in educations:
        p = doc.add_paragraph(f'• {edu}')
        p_run = p.runs[0]
        p_run.font.bold = True
        doc.add_paragraph(f'  {org} ({period}, {hours})', style='List Bullet 2')
    
    doc.add_heading('자격증', 2)
    certifications = [
        ('정보처리기사', '2025.09'),
        ('RFID-GL', '2013.11'),
        ('SCJP', '2010.04'),
        ('전기공사', '2004.08')
    ]
    
    for cert, date in certifications:
        doc.add_paragraph(f'✓ {cert} ({date})', style='List Bullet')
    
    doc.add_heading('해외 경험', 2)
    doc.add_paragraph('산업인력공단 월드잡 연수 프로그램 (2010.07 - 2011.05)', style='List Bullet')
    doc.add_paragraph('• Canadagate IT 비즈니스 실무 과정', style='List Bullet 2')
    doc.add_paragraph('• Advanced 과정(Toefl) 수업 약 4개월 수강', style='List Bullet 2')
    doc.add_paragraph('• 미국, 캐나다 Brain-based Speed Reading 세미나 참석', style='List Bullet 2')
    
    doc.add_page_break()
    
    # 6. Core Competencies 섹션
    doc.add_heading('Core Competencies', 1)
    comp_heading_run = doc.paragraphs[-1].runs[0]
    comp_heading_run.font.color.rgb = primary_color
    
    doc.add_heading('기술 역량', 2)
    doc.add_paragraph('Full-Stack Development:', style='List Bullet')
    doc.add_paragraph('• Backend: Java, Spring Framework, Spring Boot, Spring AI', style='List Bullet 2')
    doc.add_paragraph('• Frontend: HTML5/CSS3, JavaScript, jQuery, Bootstrap', style='List Bullet 2')
    doc.add_paragraph('• Mobile: Android (Java/Kotlin), iOS (Swift/SwiftUI), Flutter', style='List Bullet 2')
    doc.add_paragraph('• Database: Oracle, SQLite', style='List Bullet 2')
    doc.add_paragraph('• DevOps: Git, Docker, CI/CD (Jenkins)', style='List Bullet 2')
    
    doc.add_heading('프로젝트 경험', 2)
    doc.add_paragraph('금융권 프로젝트:', style='List Bullet')
    doc.add_paragraph('• 우리은행, 신한은행, KB국민카드/은행, 하나은행 등', style='List Bullet 2')
    doc.add_paragraph('• 안드로이드 네이티브 앱 개발', style='List Bullet 2')
    doc.add_paragraph('• 웹 서비스 전환 및 고도화', style='List Bullet 2')
    doc.add_paragraph('• 보안 및 인증 시스템 구현', style='List Bullet 2')
    
    doc.add_paragraph('기타 프로젝트:', style='List Bullet')
    doc.add_paragraph('• O2O 플랫폼 개발', style='List Bullet 2')
    doc.add_paragraph('• 증권사 MTS 개발', style='List Bullet 2')
    doc.add_paragraph('• 도시가스 검침 시스템 개발', style='List Bullet 2')
    doc.add_paragraph('• 통합 생산성 앱 개발', style='List Bullet 2')
    
    doc.add_heading('주요 강점', 2)
    strengths = [
        '금융권 프로젝트 다수 경험',
        '풀스택 개발 역량',
        '크로스 플랫폼 개발 경험',
        '최신 기술 학습 및 적용 능력'
    ]
    for strength in strengths:
        doc.add_paragraph(f'✓ {strength}', style='List Bullet')
    
    # 7. Contact 섹션
    doc.add_page_break()
    doc.add_heading('Contact', 1)
    contact_heading_run = doc.paragraphs[-1].runs[0]
    contact_heading_run.font.color.rgb = primary_color
    contact_heading_run.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('문의사항이 있으시면 언제든지 연락주세요.')
    doc.add_paragraph()
    
    doc.add_paragraph('LinkedIn: linkedin.com/in/namil-kim-a59951123')
    doc.add_paragraph('GitHub: github.com/NAM-IL')
    
    # 파일 저장 (doc 폴더에 저장)
    filename = DOC_DIR / "PORTFOLIO_PRESENTATION.docx"
    doc.save(str(filename))
    print(f"✅ 포트폴리오 DOC가 생성되었습니다: {filename}")
    print(f"📄 총 {len(doc.paragraphs)}개의 단락이 포함되어 있습니다.")
    print(f"\n💡 index.html의 내용을 기반으로 작성되었습니다.")
    print(f"📁 저장 위치: {filename}")

if __name__ == "__main__":
    try:
        create_portfolio_doc()
    except ImportError:
        print("❌ python-docx 라이브러리가 설치되지 않았습니다.")
        print("📦 설치 방법: pip install python-docx")
    except Exception as e:
        print(f"❌ 오류 발생: {e}")
        import traceback
        traceback.print_exc()
